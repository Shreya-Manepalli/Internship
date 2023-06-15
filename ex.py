import locale
import PySimpleGUI as sg
import docx  
import openpyxl  
from openpyxl import Workbook  
from openpyxl.styles import Font  
from openpyxl.styles.borders import Border, Side  

sg.theme('Purple')
SYMBOL_UP = '▲'
SYMBOL_DOWN = '▼'

def collapse(sectionLayout, visible, key):
    return sg.pin(sg.Column(sectionLayout, visible=visible, key=key))

# ---------------MAIN LAYOUT---------------
layout = [

          [sg.Text('Thousand & Decimal separator in Word document:', font=(None, 9, 'bold'))],
          [sg.Radio('1.234.567,999', 'RADIO1', default=True, size=(19, 1), key='locale'),
          sg.Radio('1,234,567.999', 'RADIO1', default=False, size=(19, 1))],

          [sg.Text('Source for Word document:', size=(25, 1), font=(None, 9, 'bold')),
           sg.InputText(key='wordDocument'), sg.FileBrowse()],
           
          [sg.Text('Source of Excel file:', size=(25, 1), font=(None, 9, 'bold'))],
          [sg.Radio('Existed Excel file:', 'RADIO0', size=(19, 1), default=True), sg.InputText(key='excelTemplate'),
           sg.FileBrowse()],
          [sg.Radio('Create a blank Excel file', 'RADIO0', default=False, key='newWb')],

          [sg.Text('Choose any of the following options:', font=(None, 9, 'bold'))],
          [sg.Checkbox('Apply cell border (All Borders)', default=False, key='border')],
          [sg.Checkbox('Copy text format (Bold, Italic & Underline)', default=False, key='formatText')],
          [sg.Checkbox('Convert cell "-" to 0', default=False, key='hyphenToZero')],
          

          [sg.Text('Choose folder to save:', size=(25, 1), font=(None, 9, 'bold')),
           sg.InputText(key='savedFolder'), sg.FolderBrowse()],
          [sg.Text('Enter file name (Ex: file.xlsx):', size=(25, 1), font=(None, 9, 'bold')),
           sg.InputText(default_text='file.xlsx', key='savedName')],

          [sg.Button('Submit'), sg.Button('Exit')],
          [sg.Output(size=(80, 3))]
          ]
# Modify the wordToExcel function
def wordToExcel(dict):
    # Load existed Excel file or create a new blank Excel file
    if not dict['newWb']:
        wb = openpyxl.load_workbook(dict['excelTemplate'])
    else:
        wb = Workbook()

    # Load the Word file
    doc = docx.Document(dict['wordDocument'])

    if dict['locale']:
        locale.setlocale(locale.LC_NUMERIC, "en_DK.UTF-8") # 1.234.567,999
    else:
        locale.setlocale(locale.LC_NUMERIC, "en_US.UTF-8") # 1,234,567.999

        
    # Set border style for Excel cells
    thin_border = Border(left=Side(style='thin'),
                         right=Side(style='thin'),
                         top=Side(style='thin'),
                         bottom=Side(style='thin'))

    # Loop through all tables in the Word file
    for t in range(len(doc.tables)):
        # Create a new sheet for each table
        sheet = wb.create_sheet(title=f'Table {t+1}')

        # Assign the table to be copied
        tableDoc = doc.tables[t]

        # Get the size of the table
        rowNumDoc = len(tableDoc.rows)
        colNumDoc = len(tableDoc.columns)

        checkMerge = []

        # Loop through all cells in the table and copy to the sheet
        for i in range(rowNumDoc):
            for j in range(colNumDoc):
                # Excel cell starts at (1, 1)
                excelCell = sheet.cell(row=i + 1, column=j + 1)

                # Word cell starts at (0, 0) = top left corner of table
                # Convert number from Word to number in Excel
                if tableDoc.cell(i, j)._tc not in checkMerge:
                    try:
                        excelCell.value = locale.atoi(
                            tableDoc.cell(i, j).text.replace('(', '-').replace(')', ''))  # integer; (4123) = -4123
                    except ValueError:
                        try:
                            excelCell.value = locale.atof(
                                tableDoc.cell(i, j).text.replace('(', '-').replace(')', ''))  # float; (4.123) = -4.123
                            
                        except ValueError:
                            if dict['hyphenToZero'] and tableDoc.cell(i, j).text == "-":
                                excelCell.value = 0
                            else:
                                excelCell.value = tableDoc.cell(i, j).text  # string
                            
                # add ._tc to the 'check merged' list
                checkMerge.append(tableDoc.cell(i, j)._tc)

                # (optional) format Excel cells
                if dict['border']:
                    excelCell.border = thin_border  # apply cell border
                if dict['formatText']:
                    if tableDoc.cell(i, j).text:  # exclude empty cells in Word
                        try:
                            checkBold = tableDoc.cell(i, j).paragraphs[0].runs[
                                0].bold  # check format in Word: True/False/None
                            checkItalic = tableDoc.cell(i, j).paragraphs[0].runs[
                                0].italic  # check format in Word: True/False/None
                            checkUnderline = tableDoc.cell(i, j).paragraphs[0].runs[
                                0].underline  # check format in Word: True/False/None
                        except:
                            continue

                        # underline: {'single', 'singleAccounting', 'doubleAccounting', 'double'}
                        excelCell.font = Font(bold=checkBold, italic=checkItalic,
                                              underline=('single' if checkUnderline else None))

                # Apply cell border
                excelCell.border = thin_border

    # Save the Excel file
    if not dict['savedName'].endswith('.xlsx'):
        dict['savedName'] += '.xlsx'
    wb.save((dict['savedFolder'] + '/' + dict['savedName']))



# load GUI
window = sg.Window('Copy Word tables to Excel', layout)

while True:  # Event Loop
    event, values = window.read()
    # print(event, values)
    if event == sg.WIN_CLOSED or event == 'Exit':
        break
    if event == 'Submit':
        # Print message to Output box
        print('Running............')
        # convert the 'window dict' to 'normal dict' because 'window dict' values have different class(es)
        keys = ['excelTemplate', 'newWb', 'wordDocument',
                'locale',
                'border','hyphenToZero', 'formatText',
                'savedFolder', 'savedName'
                ]
        convertDict = {x: window[x].get() for x in keys}

        # run 'wordToExcel'
        if convertDict['wordDocument'] and convertDict['savedFolder']:
            wordToExcel(convertDict)
        # Print message to Output box
        print('Done!!!!')

    if event.startswith('-OPEN SEC1-'):
        opened1 = not opened1
        window['-OPEN SEC1-'].update(SYMBOL_UP if opened1 else SYMBOL_DOWN)
        window['-SEC1-'].update(visible=opened1)

window.close()


