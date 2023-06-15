import locale
import PySimpleGUI as sg
import docx  
import openpyxl  
from openpyxl import Workbook  
from openpyxl.styles import Font  
from openpyxl.styles.borders import Border, Side  

sg.theme('Purple')
SYMBOL_UP = '▲'
SYMBOL_DOWN = '▼'

def collapse(sectionLayout, visible, key):
    return sg.pin(sg.Column(sectionLayout, visible=visible, key=key))

# LAYOUT
layout = [

          [sg.Text('Thousand & Decimal separator in Word document:', font=(None, 9, 'bold'))],
          [sg.Radio('1.234.567,999', 'RADIO1', default=True, size=(19, 1), key='locale'),
          sg.Radio('1,234,567.999', 'RADIO1', default=False, size=(19, 1))],

          [sg.Text('Source for Word document:', size=(25, 1), font=(None, 9, 'bold')),
           sg.InputText(key='wordDocument'), sg.FileBrowse()],
           
          [sg.Text('Source of Excel file:', size=(25, 1), font=(None, 9, 'bold'))],
          [sg.Radio('Create a blank Excel file', 'RADIO0', default=False, key='newWb')],

          [sg.Text('Choose any of the following options:', font=(None, 9, 'bold'))],
          [sg.Checkbox('Apply cell border (All Borders)', default=False, key='border')],
          [sg.Checkbox('Copy text format (Bold, Italic & Underline)', default=False, key='formatText')],
          [sg.Checkbox('Convert cell "-" to 0', default=False, key='hyphenToZero')],
          

          [sg.Text('Choose folder to save:', size=(25, 1), font=(None, 9, 'bold')),
           sg.InputText(key='savedFolder'), sg.FolderBrowse()],
          [sg.Text('Enter file name (Ex: file.xlsx):', size=(25, 1), font=(None, 9, 'bold')),
           sg.InputText(default_text='file.xlsx', key='savedName')],

          [sg.Button('Submit'), sg.Button('Exit')],
          [sg.Output(size=(80, 3))]
          ]

def wordToExcel(dict):

    if dict['newWb']:
        wb = Workbook()

    doc = docx.Document(dict['wordDocument'])

    if dict['locale']:
        locale.setlocale(locale.LC_NUMERIC, "en_DK.UTF-8") 
    else:
        locale.setlocale(locale.LC_NUMERIC, "en_US.UTF-8") 

    thin_border = Border(left=Side(style='thin'),
                         right=Side(style='thin'),
                         top=Side(style='thin'),
                         bottom=Side(style='thin'))


    for t in range(len(doc.tables)):
        sheet = wb.create_sheet(title=f'Table {t+1}')
        tableDoc = doc.tables[t]


        rowNumDoc = len(tableDoc.rows)
        colNumDoc = len(tableDoc.columns)

        checkMerge = []

        for i in range(rowNumDoc):
            for j in range(colNumDoc):
 
                excelCell = sheet.cell(row=i + 1, column=j + 1)

                if tableDoc.cell(i, j)._tc not in checkMerge:
                    try:
                        excelCell.value = locale.atoi(
                            tableDoc.cell(i, j).text.replace('(', '-').replace(')', ''))  
                    except ValueError:
                        try:
                            excelCell.value = locale.atof(
                                tableDoc.cell(i, j).text.replace('(', '-').replace(')', ''))  
                            
                        except ValueError:
                            if dict['hyphenToZero'] and tableDoc.cell(i, j).text == "-":
                                excelCell.value = 0
                            else:
                                excelCell.value = tableDoc.cell(i, j).text  
                            

                checkMerge.append(tableDoc.cell(i, j)._tc)


                if dict['border']:
                    excelCell.border = thin_border  
                if dict['formatText']:
                    if tableDoc.cell(i, j).text:  
                        try:
                            checkBold = tableDoc.cell(i, j).paragraphs[0].runs[
                                0].bold  
                            checkItalic = tableDoc.cell(i, j).paragraphs[0].runs[
                                0].italic  
                            checkUnderline = tableDoc.cell(i, j).paragraphs[0].runs[
                                0].underline  
                        except:
                            continue


                        excelCell.font = Font(bold=checkBold, italic=checkItalic,
                                              underline=('single' if checkUnderline else None))
                    excelCell.border = thin_border


    if not dict['savedName'].endswith('.xlsx'):
        dict['savedName'] += '.xlsx'
    wb.save((dict['savedFolder'] + '/' + dict['savedName']))


window = sg.Window('Copy Word tables to Excel', layout)

while True:  
    event, values = window.read()

    if event == sg.WIN_CLOSED or event == 'Exit':
        break
    if event == 'Submit':

        print('Running............')

        keys = ['newWb', 'wordDocument',
                'locale',
                'border','hyphenToZero', 'formatText',
                'savedFolder', 'savedName'
                ]
        convertDict = {x: window[x].get() for x in keys}

        if convertDict['wordDocument'] and convertDict['savedFolder']:
            wordToExcel(convertDict)
        print('Done!!!!')

    if event.startswith('-OPEN SEC1-'):
        opened1 = not opened1
        window['-OPEN SEC1-'].update(SYMBOL_UP if opened1 else SYMBOL_DOWN)
        window['-SEC1-'].update(visible=opened1)

window.close()
