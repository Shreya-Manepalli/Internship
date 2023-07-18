
import re
import PyPDF2
import docx2txt
import locale
import PySimpleGUI as sg
import docx
from openpyxl import Workbook  
from openpyxl.styles import Font  
from openpyxl.styles.borders import Border, Side 
import tabula
import pandas as pd 
from docx import Document

sg.theme('Purple')
SYMBOL_UP = '▲'
SYMBOL_DOWN = '▼'

def collapse(sectionLayout, visible, key):
    return sg.pin(sg.Column(sectionLayout, visible=visible, key=key))

# LAYOUT
layout = [

          [sg.Text('Thousand & Decimal separator in Word document:', font=(None, 9, 'bold'))],
          [sg.Radio('1.234.567,999', 'RADIO1', default=True, size=(19, 1), key='locale'),
          sg.Radio('1,234,567.999', 'RADIO1', default=False, size=(19, 1))],

          [sg.Text('Source for Word document:', size=(25, 1), font=(None, 9, 'bold')),
           sg.InputText(key='wordDocument'), sg.FileBrowse()],

           [sg.Text('Source for PDF document:', size=(25, 1), font=(None, 9, 'bold')),
            sg.InputText(key='pdfDocument'), sg.FileBrowse()],

          [sg.Text('Source of Excel file:', size=(25, 1), font=(None, 9, 'bold'))],
          [sg.Radio('Create a blank Excel file', 'RADIO0', default=False, key='newWb')],

          [sg.Text('Choose any of the following options:', font=(None, 9, 'bold'))],
          [sg.Checkbox('Apply cell border (All Borders)', default=False, key='border')],
          [sg.Checkbox('Copy text format (Bold, Italic & Underline)', default=False, key='formatText')],
          [sg.Checkbox('Convert cell "-" to 0', default=False, key='hyphenToZero')],

          [sg.Text('Choose folder to save the word tables:', size=(25, 1), font=(None, 9, 'bold')),
           sg.InputText(key='savedFolder1'), sg.FolderBrowse()],
          [sg.Text('Enter file name (Ex: file.xlsx):', size=(25, 1), font=(None, 9, 'bold')),
           sg.InputText(default_text='file1.xlsx', key='savedName1')],

           [sg.Text('Choose folder to save the pdf tables:', size=(25, 1), font=(None, 9, 'bold')),
            sg.InputText(key='savedFolder2'), sg.FolderBrowse()],
            [sg.Text('Enter file name (Ex: file.xlsx):', size=(25, 1), font=(None, 9, 'bold')),
            sg.InputText(default_text='file2.xlsx', key='savedName2')],

          [sg.Text('Extract Relevant Text and Equations from the file:', font=(None, 9, 'bold'))],

          [sg.Text('Enter the word(s) or phrase(s) to search for (separated by a semicolon):', font=(None, 9, 'bold'))],
          [sg.InputText(key='searchWord')],
          [sg.Text('File Path (PDF or DOCX):', font=(None, 9, 'bold'))],
          [sg.InputText(key='filePath'), sg.FileBrowse()],
          [sg.Text('Output File Path (DOCX):', font=(None, 9, 'bold'))],
          [sg.InputText(key='OutputfilePath'), sg.FileBrowse()],

          [sg.Button('Extract Equations'),sg.Button('Extract Sentences')], 
          [sg.Button('Submit'), sg.Button('Exit')],
          [sg.Output(size=(80, 10))]

         ]


# Function to extract sentences containing a specific word from a PDF file
def save_sentences_to_word_file(sentences, output_filename):
    doc = docx.Document()
    for word, sentences in sentences.items():
        doc.add_paragraph(f"The relevant text for the word '{word}' is:")
        for sentence in sentences:
            doc.add_paragraph(sentence)
        doc.add_paragraph("\n")
    doc.save(output_filename)

def extract_sentences_with_word_from_pdf(filename, words):
    sentences = {word: [] for word in words}
    previous_sentence = None
    sentence_set = set()
    stop_extraction = False

    with open(filename, "rb") as file:
        reader = PyPDF2.PdfReader(file)
        num_pages = len(reader.pages)

        for page_num in range(num_pages):
            page = reader.pages[page_num]
            text = page.extract_text()

            if stop_extraction:
                break

            for word in words:
                if " " in word:
                    found_sentences = re.findall(r"([^.!?]*{}[^.!?]*[.!?])".format(re.escape(word)), text)
                else:
                    found_sentences = re.findall(r"([^.!?]*\b{}\b[^.!?]*[.!?])".format(re.escape(word)), text)

                for sentence in found_sentences:
                    if sentence in sentence_set:
                        continue  # Skip repeated sentence

                    if previous_sentence is not None and previous_sentence == sentence:
                        continue  # Skip repeated sentence

                    if sentence.startswith("Figure") or sentence.startswith("Page"):
                        continue  # Skip sentences starting with "Figure" or "Page"

                    if "references" in sentence.lower():
                        stop_extraction = True
                        break

                    sentences[word].append(sentence)
                    sentence_set.add(sentence)
                    previous_sentence = sentence

    return sentences

def extract_sentences_with_word_from_docx(filename, words):
    sentences = {word: [] for word in words}
    previous_sentence = None
    sentence_set = set()
    stop_extraction = False

    text = docx2txt.process(filename)

    for word in words:
        if " " in word:
            found_sentences = re.findall(r"([^.!?]*{}[^.!?]*[.!?])".format(re.escape(word)), text)
        else:
            found_sentences = re.findall(r"([^.!?]*\b{}\b[^.!?]*[.!?])".format(re.escape(word)), text)

        for sentence in found_sentences:
            if sentence in sentence_set:
                continue  # Skip repeated sentence

            if previous_sentence is not None and previous_sentence == sentence:
                continue  # Skip repeated sentence

            if sentence.startswith("Figure") or sentence.startswith("Page"):
                continue  # Skip sentences starting with "Figure" or "Page"

            if "references" in sentence.lower():
                stop_extraction = True
                break

            sentences[word].append(sentence)
            sentence_set.add(sentence)
            previous_sentence = sentence

    return sentences

def extract_equations_from_pdf(pdf_path):
    equations = []
    
    with open(pdf_path, 'rb') as file:
        pdf_reader = PyPDF2.PdfReader(file)
        num_pages = len(pdf_reader.pages)
        
        for page_number in range(num_pages):
            page = pdf_reader.pages[page_number]
            text = page.extract_text()
            equation_pattern = r'[^\n=]+=[^\n]+'
            matches = re.findall(equation_pattern, text)
            equations.extend(matches)
    
    return equations

def extract_equations_from_docx(docx_path):
    equations = []

    text = docx2txt.process(docx_path)
    equation_pattern = r'[^\n=]+=[^\n]+'
    matches = re.findall(equation_pattern, text)
    equations.extend(matches)
    
    return equations

def save_equations_to_word_file(equations, output_filename):
    doc = docx.Document()
    for equation in equations:
        doc.add_paragraph(equation)
    doc.save(output_filename)

    
def wordToExcel(dict):

    if dict['newWb']:
        wb = Workbook()

    doc = docx.Document(dict['wordDocument'])

    if dict['locale']:
        locale.setlocale(locale.LC_NUMERIC, "en_DK.UTF-8") 
    else:
        locale.setlocale(locale.LC_NUMERIC, "en_US.UTF-8") 

    thin_border = Border(left=Side(style='thin'),
                         right=Side(style='thin'),
                         top=Side(style='thin'),
                         bottom=Side(style='thin'))


    for t in range(len(doc.tables)):
        sheet = wb.create_sheet(title=f'Table {t+1}')
        tableDoc = doc.tables[t]


        rowNumDoc = len(tableDoc.rows)
        colNumDoc = len(tableDoc.columns)

        checkMerge = []

        for i in range(rowNumDoc):
            for j in range(colNumDoc):
 
                excelCell = sheet.cell(row=i + 1, column=j + 1)

                if tableDoc.cell(i, j)._tc not in checkMerge:
                    try:
                        excelCell.value = locale.atoi(
                            tableDoc.cell(i, j).text.replace('(', '-').replace(')', ''))  
                    except ValueError:
                        try:
                            excelCell.value = locale.atof(
                                tableDoc.cell(i, j).text.replace('(', '-').replace(')', ''))  
                            
                        except ValueError:
                            if dict['hyphenToZero'] and tableDoc.cell(i, j).text == "-":
                                excelCell.value = 0
                            else:
                                excelCell.value = tableDoc.cell(i, j).text  
                            

                checkMerge.append(tableDoc.cell(i, j)._tc)


                if dict['border']:
                    excelCell.border = thin_border  
                if dict['formatText']:
                    if tableDoc.cell(i, j).text:  
                        try:
                            checkBold = tableDoc.cell(i, j).paragraphs[0].runs[
                                0].bold  
                            checkItalic = tableDoc.cell(i, j).paragraphs[0].runs[
                                0].italic  
                            checkUnderline = tableDoc.cell(i, j).paragraphs[0].runs[
                                0].underline  
                        except:
                            continue


                        excelCell.font = Font(bold=checkBold, italic=checkItalic,
                                              underline=('single' if checkUnderline else None))
                    excelCell.border = thin_border


    if not dict['savedName1'].endswith('.xlsx'):
        dict['savedName1'] += '.xlsx'
    wb.save((dict['savedFolder1'] + '/' + dict['savedName1']))


##################################################################


def pdfToExcel(dict):
    if dict['newWb']:
        wb = Workbook()
    tables = tabula.read_pdf(dict['pdfDocument'], pages='all', multiple_tables=True)

    if dict['locale']:
        locale.setlocale(locale.LC_NUMERIC, "en_DK.UTF-8")  
    else:
        locale.setlocale(locale.LC_NUMERIC, "en_US.UTF-8")  

    
    thin_border = Border(left=Side(style='thin'),
                         right=Side(style='thin'),
                         top=Side(style='thin'),
                         bottom=Side(style='thin'))

    
    for table_num, table in enumerate(tables):
        
        sheet = wb.create_sheet(title=f'Table {table_num + 1}')
        excel_writer = pd.ExcelWriter((dict['savedFolder2'] + '/' + dict['savedName2']), engine='xlsxwriter')

    for i, table in enumerate(tables):
    
        sheet_name = f"Table {i+1}"
        table.to_excel(excel_writer, sheet_name=sheet_name, index=False)

    excel_writer.save()

        
    for i, row in enumerate(tables):
            
            for j, cell in enumerate(row):
                
                excelCell = sheet.cell(row=i + 1, column=j + 1)

                
                try:
                    excelCell.value = locale.atof(cell.replace('(', '-').replace(')', ''))
                except ValueError:
                    if dict['hyphenToZero'] and cell == "-":
                        excelCell.value = 0
                    else:
                        excelCell.value = cell

                
                if dict['border']:
                    excelCell.border = thin_border

                
                if dict['formatText']:
                    excelCell.font = Font(bold=False,
                                          italic=False,
                                          underline=False)

    
    if not dict['savedName2'].endswith('.xlsx'):
        dict['savedName2'] += '.xlsx'
    wb.save((dict['savedFolder2'] + '/' + dict['savedName2']))

window = sg.Window('Copy Word Tables or PDF Tables to Excel', layout)

while True:  
    event, values = window.read()

    if event == sg.WIN_CLOSED or event == 'Exit':
        break
    if event == 'Submit':

        print('Running............')

        keys = ['newWb', 'wordDocument', 'pdfDocument', 'locale',
                'border', 'hyphenToZero', 'formatText',
                'savedFolder1', 'savedName1',
                'savedFolder2', 'savedName2'
        ]
        convertDict = {x: window[x].get() for x in keys}

        if convertDict['wordDocument'] and convertDict['savedFolder1']:
            wordToExcel(convertDict)

        if convertDict['pdfDocument'] and convertDict['savedFolder2']:
            pdfToExcel(convertDict)
        print('Done!!!!')

    if event == 'Extract Sentences':
        search_word = values['searchWord']
        file_path = values['filePath']
        output_filename = values['OutputfilePath']
        words = search_word.split(";")
        words = [word.strip() for word in words]
        sentences = {}

        if file_path.lower().endswith('.pdf'):
            sentences = extract_sentences_with_word_from_pdf(file_path, words)
        elif file_path.lower().endswith('.docx'):
            sentences = extract_sentences_with_word_from_docx(file_path, words)
        else:
            print("Unsupported file format.")

        if sentences:
            save_sentences_to_word_file(sentences, output_filename)
        else:
            print("No relevant text found for the given word(s) or phrase(s). Please try again.")

    if event == 'Extract Equations':
        file_path = values['filePath']
        output_filename = values['OutputfilePath']
        equations = []

        if file_path.lower().endswith('.pdf'):
            equations = extract_equations_from_pdf(file_path)
        elif file_path.lower().endswith('.docx'):
            equations = extract_equations_from_docx(file_path)
        else:
            print("Unsupported file format.")
        save_equations_to_word_file(equations, output_filename)

window.close()