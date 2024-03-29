#!/usr/bin/env python



'''
Description: This code does numerous functions. From a pdf/word file, it is possible to extract tables, equations, table captions, figure captions, and images.
It produces the desired output from a word or pdf file that is provided as input.

'''



__author__ = "Shreya Manepalli"
__version__ = "1.1.0"
__email__ = "Shreya.Manepalli@gknaerospace.com"




############################################################# CODE #############################################################


#IMPORTING LIBRARIES
import re
import PyPDF2
import docx2txt
import locale
import PySimpleGUI as sg
import docx
from openpyxl import Workbook  
from openpyxl.styles import Font  
from openpyxl.styles.borders import Border, Side 
import tabula
import pandas as pd 
from docx import Document
from PyPDF2 import PdfReader
import fitz
import os
from PIL import Image


########################### GUI ###################################


sg.theme('Purple')
SYMBOL_UP = '▲'
SYMBOL_DOWN = '▼'

def collapse(sectionLayout, visible, key):
    return sg.pin(sg.Column(sectionLayout, visible=visible, key=key))

# LAYOUT
layout = [

          [sg.Text('Thousand & Decimal separator in Word document:', font=(None, 9, 'bold'))],
          [sg.Radio('1.234.567,999', 'RADIO1', default=True, size=(19, 1), key='locale'),
          sg.Radio('1,234,567.999', 'RADIO1', default=False, size=(19, 1))],

          [sg.Text('Source for Word document:', size=(47, 1), font=(None, 9, 'bold')),
           sg.InputText(key='wordDocument'), sg.FileBrowse()],

           [sg.Text('Source for PDF document:', size=(47, 1), font=(None, 9, 'bold')),
            sg.InputText(key='pdfDocument'), sg.FileBrowse()],

          [sg.Text('Source of Excel file:', size=(47, 1), font=(None, 9, 'bold')),
          sg.Radio('Create a blank Excel file', 'RADIO0', default=False, key='newWb')],

          [sg.Text('Choose any of the following options:', font=(None, 9, 'bold'))],
          [sg.Checkbox('Apply cell border (All Borders)', default=False, key='border')],
          [sg.Checkbox('Copy text format (Bold, Italic & Underline)', default=False, key='formatText')],
          [sg.Checkbox('Convert cell "-" to 0', default=False, key='hyphenToZero')],

          [sg.Text('Choose folder to save the word tables:', size=(47, 1), font=(None, 9, 'bold')),
           sg.InputText(key='savedFolder1'), sg.FolderBrowse()],
          [sg.Text('Enter file name (Ex: file.xlsx):', size=(47, 1), font=(None, 9, 'bold')),
           sg.InputText(key='savedName1')],

           [sg.Text('Choose folder to save the pdf tables:', size=(47, 1), font=(None, 9, 'bold')),
            sg.InputText(key='savedFolder2'), sg.FolderBrowse()],
            [sg.Text('Enter file name (Ex: file.xlsx):', size=(47, 1), font=(None, 9, 'bold')),
            sg.InputText(key='savedName2')],
          [sg.Button('Extract Tables')],

          [sg.Text('Source for PDF or Word Document (PDF or DOCX):', size=(47, 1), font=(None, 9, 'bold')),
          sg.InputText(key='filePath'), sg.FileBrowse()],

          [sg.Text('Enter the word(s) or phrase(s)(separated by a semicolon):', size=(47, 1), font=(None, 9, 'bold')),
          sg.InputText(key='searchWord')],
          [sg.Text('Choose file to save relevant text(DOCX file):', size=(47, 1), font=(None, 9, 'bold')),
          sg.InputText(key='OutputfilePath'), sg.FileBrowse()],
          [sg.Button('Extract Relevant Text')], 

          [sg.Text('Choose file to save equations (DOCX file):', size=(47, 1), font=(None, 9, 'bold')),
          sg.InputText(key='OutputfilePath1'), sg.FileBrowse()],
          [sg.Button('Extract Equations')],

          [sg.Text('Choose file to save the figure captions (DOCX file):', size=(47, 1), font=(None, 9, 'bold')),
          sg.InputText(key='OutputfilePath2'), sg.FileBrowse()],
          [sg.Button('Extract Figure Captions')],
          [sg.Text('Choose file to save the table captions (DOCX file):', size=(47, 1), font=(None, 9, 'bold')),
          sg.InputText(key='OutputfilePath3'), sg.FileBrowse()],
		  [sg.Button('Extract Table captions')],
          [sg.Text('Choose folder to save the extracted images: ', size=(47, 1), font=(None, 9, 'bold')),
          sg.InputText(key='Outputfolderpath'), sg.FolderBrowse()],         
          [sg.Button('Extract Images')],
          [sg.Output(size=(100, 7))],
          [sg.Button('Exit')],

         ]


########################### EXTRACT RELEVANT TEXT BASED ON KEYWORD INPUT #################################


#Function to extract relevant text based on the keyword input from pdf files
def extract_sentences_with_word_from_pdf(filename, words):
    sentences = {word: [] for word in words}
    previous_sentence = None
    sentence_set = set()
    stop_extraction = False

    with open(filename, "rb") as file:
        reader = PyPDF2.PdfReader(file)
        num_pages = len(reader.pages)

        for page_num in range(num_pages):
            page = reader.pages[page_num]
            text = page.extract_text()

            if stop_extraction:
                break

            for word in words:
                if " " in word:
                    found_sentences = re.findall(r"([^.!?]*{}[^.!?]*[.!?])".format(re.escape(word)), text)
                else:
                    found_sentences = re.findall(r"([^.!?]*\b{}\b[^.!?]*[.!?])".format(re.escape(word)), text)

                for sentence in found_sentences:
                    if sentence in sentence_set:
                        continue  # Skip repeated sentence

                    if previous_sentence is not None and previous_sentence == sentence:
                        continue  # Skip repeated sentence

                    if sentence.startswith("Figure") or sentence.startswith("Page"):
                        continue  # Skip sentences starting with "Figure" or "Page"

                    if "references" in sentence.lower():
                        stop_extraction = True
                        break

                    sentences[word].append(sentence)
                    sentence_set.add(sentence)
                    previous_sentence = sentence

    return sentences


#Function to extract relevant text based on the keyword input from word files
def extract_sentences_with_word_from_docx(filename, words):
    sentences = {word: [] for word in words}
    previous_sentence = None
    sentence_set = set()
    stop_extraction = False

    text = docx2txt.process(filename)

    for word in words:
        if " " in word:
            found_sentences = re.findall(r"([^.!?]*{}[^.!?]*[.!?])".format(re.escape(word)), text)
        else:
            found_sentences = re.findall(r"([^.!?]*\b{}\b[^.!?]*[.!?])".format(re.escape(word)), text)

        for sentence in found_sentences:
            if sentence in sentence_set:
                continue  # Skip repeated sentence

            if previous_sentence is not None and previous_sentence == sentence:
                continue  # Skip repeated sentence

            if sentence.startswith("Figure") or sentence.startswith("Page"):
                continue  # Skip sentences starting with "Figure" or "Page"

            if "references" in sentence.lower():
                stop_extraction = True
                break

            sentences[word].append(sentence)
            sentence_set.add(sentence)
            previous_sentence = sentence

    return sentences


# Function to save the extracted Relevant Text containing a specific word into a word file
def save_relevanttext_to_word_file(sentences, output_filename):
    doc = docx.Document()
    for word, sentences in sentences.items():
        doc.add_paragraph(f"The relevant text for the word '{word}' is:")
        for sentence in sentences:
            doc.add_paragraph(sentence)
        doc.add_paragraph("\n")
    doc.save(output_filename)


########################### EXTRACT EQUATIONS #################################


#Function to extract all the equations from pdf files
def extract_equations_from_pdf(pdf_path):
    equations = []
    
    with open(pdf_path, 'rb') as file:
        pdf_reader = PyPDF2.PdfReader(file)
        num_pages = len(pdf_reader.pages)
        
        for page_number in range(num_pages):
            page = pdf_reader.pages[page_number]
            text = page.extract_text()
            equation_pattern = r'[^\n=]+=[^\n]+'
            matches = re.findall(equation_pattern, text)
            equations.extend(matches)
    
    return equations


#Function to extract all the equations from word files
def extract_equations_from_docx(docx_path):
    equations = []

    text = docx2txt.process(docx_path)
    equation_pattern = r'[^\n=]+=[^\n]+'
    matches = re.findall(equation_pattern, text)
    equations.extend(matches)
    
    return equations


#Function to save the extracted equations into a word file
def save_equations_to_word_file(equations, output_filename):
    doc = docx.Document()
    for equation in equations:
        doc.add_paragraph(equation)
    doc.save(output_filename)


########################### EXTRACT FIGURE CAPTIONS #################################


#Function to extract all the figure captions from pdf files
def extract_figure_sentences_from_pdf(pdf_path):
    # Read the PDF file
    with open(pdf_path, 'rb') as file:
        pdf = PdfReader(file)
        num_pages = len(pdf.pages)

        figure_sentences = []
        for page_num in range(num_pages):
            page = pdf.pages[page_num]
            text = page.extract_text()

            # Extract sentences starting with "Fig" or "Figure"
            sentences = re.findall(r'(?i)((?:Fig|Figure).*?)(?=\n)', text)
            figure_sentences.extend(sentences)

    return figure_sentences


#Function to extract all the figure captions from word files
def extract_figure_sentences_from_word(docx_path):
    doc = docx.Document(docx_path)
    figure_sentences = []
    for paragraph in doc.paragraphs:
        # Extract sentences starting with "Fig" or "Figure"
        sentences = re.findall(r'(?i)((?:Fig|Figure).*?)(?=\n|$)', paragraph.text)
        figure_sentences.extend(sentences)
    return figure_sentences


#Function to save the extracted figure captions into a word file
def save_figcaptions_to_word(figure_sentences, output_path):
    doc = docx.Document()

    # Add each sentence as a paragraph in the Word document
    for figure_sentence in figure_sentences:
        doc.add_paragraph(figure_sentence)

    # Save the Word document
    doc.save(output_path)


########################### EXTRACT TABLE CAPTIONS #################################


#Function to extract all the table captions from pdf files
def extract_table_sentences_from_pdf(pdf_path):
    # Read the PDF file
    with open(pdf_path, 'rb') as file:
        pdf = PdfReader(file)
        num_pages = len(pdf.pages)

        table_sentences = []
        table_seen = {}  # Dictionary to keep track of seen Table sentences
        for page_num in range(num_pages):
            page = pdf.pages[page_num]
            text = page.extract_text()

            # Extract sentences starting with "Table" followed by a number and a dot
            sentences = re.findall(r'(?i)(Table\s+\d+\s*\..*?)(?=\n)', text)
            for sentence in sentences:
                # Check if the Table sentence is not seen before
                if sentence not in table_seen:
                    table_seen[sentence] = True
                    table_sentences.append(sentence)

    return table_sentences


#Function to extract all the table captions from word files
def extract_table_sentences_from_word(docx_path):
    doc = docx.Document(docx_path)
    table_sentences = []
    table_seen = {}  # Dictionary to keep track of seen Table sentences
    for paragraph in doc.paragraphs:
        # Extract sentences starting with "Table" followed by a number and a dot
        sentences = re.findall(r'(?i)(Table\s+\d+\s*\..*?)(?=\n|$)', paragraph.text)
        for sentence in sentences:
            # Check if the Table sentence is not seen before
            if sentence not in table_seen:
                table_seen[sentence] = True
                table_sentences.append(sentence)
    return table_sentences


#Function to save the extracted table captions into a word file
def save_tablecaptions_to_word(sentences, output_path):
    doc = docx.Document()

    # Add each sentence as a paragraph in the Word document
    for sentence in sentences:
        doc.add_paragraph(sentence)

    # Save the Word document
    doc.save(output_path)


########################### EXTRACT IMAGES #################################


#Function to extract all the images from pdf files
def extract_images_from_pdf(file_path):
    pdf_file = fitz.open(file_path)
    page_nums = len(pdf_file)

    images_list = []
    image_count = 0

    for page_num in range(page_nums):
        page_content = pdf_file[page_num]
        images_list.extend(page_content.get_images())

    if len(images_list) == 0:
        raise ValueError(f'No images found in {file_path}')

    for i, img in enumerate(images_list, start=1):
        xref = img[0]
        base_image = pdf_file.extract_image(xref)
        image_bytes = base_image['image']
        image_ext = base_image['ext']
        image_name = str(image_count) + '.' + image_ext

        image_count += 1
        yield image_name, image_bytes
    
    pdf_file.close()


#Function to extract all the images from word files
def extract_images_from_word(file_path):
    doc = docx.Document(file_path)

    images_list = []
    image_count = 0

    for rel in doc.part.rels:
        for rId, relObj in enumerate(doc.part.rels[rel].target_part.rels):
            if "image" in doc.part.rels[rel].target_part.rels[relObj].target_ref:
                image = doc.part.rels[rel].target_part.rels[relObj].target_part.blob
                images_list.append(image)
                image_count += 1

                image_name = str(image_count) + '.png'  # Save as PNG format, but you can change the format if needed.

                yield image_name, image

    if len(images_list) == 0:
        raise ValueError(f'No images found in {file_path}')
    

#Function to save the extracted images into a folder
def save_images_to_folders(images, output_folder):
    os.makedirs(output_folder, exist_ok=True)

    for image_name, image_bytes in images:
        image_path = os.path.join(output_folder, image_name)
        with open(image_path, 'wb') as image_file:
            image_file.write(image_bytes)


########################### EXTRACT TABLES FROM WORD DOCUMENTS#################################


#Function to extract all the tables from word files
def wordToExcel(dict):

    if dict['newWb']:
        wb = Workbook()

    doc = docx.Document(dict['wordDocument'])

    if dict['locale']:
        locale.setlocale(locale.LC_NUMERIC, "en_DK.UTF-8") 
    else:
        locale.setlocale(locale.LC_NUMERIC, "en_US.UTF-8") 

    thin_border = Border(left=Side(style='thin'),
                         right=Side(style='thin'),
                         top=Side(style='thin'),
                         bottom=Side(style='thin'))


    for t in range(len(doc.tables)):
        sheet = wb.create_sheet(title=f'Table {t+1}')
        tableDoc = doc.tables[t]


        rowNumDoc = len(tableDoc.rows)
        colNumDoc = len(tableDoc.columns)

        checkMerge = []

        for i in range(rowNumDoc):
            for j in range(colNumDoc):
 
                excelCell = sheet.cell(row=i + 1, column=j + 1)

                if tableDoc.cell(i, j)._tc not in checkMerge:
                    try:
                        excelCell.value = locale.atoi(
                            tableDoc.cell(i, j).text.replace('(', '-').replace(')', ''))  
                    except ValueError:
                        try:
                            excelCell.value = locale.atof(
                                tableDoc.cell(i, j).text.replace('(', '-').replace(')', ''))  
                            
                        except ValueError:
                            if dict['hyphenToZero'] and tableDoc.cell(i, j).text == "-":
                                excelCell.value = 0
                            else:
                                excelCell.value = tableDoc.cell(i, j).text  
                            

                checkMerge.append(tableDoc.cell(i, j)._tc)


                if dict['border']:
                    excelCell.border = thin_border  
                if dict['formatText']:
                    if tableDoc.cell(i, j).text:  
                        try:
                            checkBold = tableDoc.cell(i, j).paragraphs[0].runs[
                                0].bold  
                            checkItalic = tableDoc.cell(i, j).paragraphs[0].runs[
                                0].italic  
                            checkUnderline = tableDoc.cell(i, j).paragraphs[0].runs[
                                0].underline  
                        except:
                            continue


                        excelCell.font = Font(bold=checkBold, italic=checkItalic,
                                              underline=('single' if checkUnderline else None))
                    excelCell.border = thin_border


    if not dict['savedName1'].endswith('.xlsx'):
        dict['savedName1'] += '.xlsx'
    wb.save((dict['savedFolder1'] + '/' + dict['savedName1']))


########################### EXTRACT TABLES FROM PDF FILES#################################


#Function to extract all the tables from pdf files
def pdfToExcel(dict):
    if dict['newWb']:
        wb = Workbook()
    tables = tabula.read_pdf(dict['pdfDocument'], pages='all', multiple_tables=True)

    if dict['locale']:
        locale.setlocale(locale.LC_NUMERIC, "en_DK.UTF-8")  
    else:
        locale.setlocale(locale.LC_NUMERIC, "en_US.UTF-8")  

    
    thin_border = Border(left=Side(style='thin'),
                         right=Side(style='thin'),
                         top=Side(style='thin'),
                         bottom=Side(style='thin'))

    
    for table_num, table in enumerate(tables):
        
        sheet = wb.create_sheet(title=f'Table {table_num + 1}')
        excel_writer = pd.ExcelWriter((dict['savedFolder2'] + '/' + dict['savedName2']), engine='xlsxwriter')

    for i, table in enumerate(tables):
    
        sheet_name = f"Table {i+1}"
        table.to_excel(excel_writer, sheet_name=sheet_name, index=False)

    excel_writer.save()

        
    for i, row in enumerate(tables):
            
            for j, cell in enumerate(row):
                
                excelCell = sheet.cell(row=i + 1, column=j + 1)

                
                try:
                    excelCell.value = locale.atof(cell.replace('(', '-').replace(')', ''))
                except ValueError:
                    if dict['hyphenToZero'] and cell == "-":
                        excelCell.value = 0
                    else:
                        excelCell.value = cell

                
                if dict['border']:
                    excelCell.border = thin_border

                
                if dict['formatText']:
                    excelCell.font = Font(bold=False,
                                          italic=False,
                                          underline=False)

    
    if not dict['savedName2'].endswith('.xlsx'):
        dict['savedName2'] += '.xlsx'
    wb.save((dict['savedFolder2'] + '/' + dict['savedName2']))


###########################LINKING GUI TO THE CODE#################################


#Connects the code to the GUI
window = sg.Window('Data Extractor', layout)

while True:  
    event, values = window.read()

    if event == 'Extract Tables':

        print('Extracting Tables............')

        keys = ['newWb', 'wordDocument', 'pdfDocument', 'locale',
                'border', 'hyphenToZero', 'formatText',
                'savedFolder1', 'savedName1',
                'savedFolder2', 'savedName2'
        ]
        convertDict = {x: window[x].get() for x in keys}

        if convertDict['wordDocument'] and convertDict['savedFolder1']:
            wordToExcel(convertDict)
            print('Done, Extracted Tables!!!!')
        if convertDict['pdfDocument'] and convertDict['savedFolder2']:
            pdfToExcel(convertDict)
            print('Done, Extracted Tables!!!!')


    if event == 'Extract Relevant Text':

        print('Extracting Relevent text............')

        search_word = values['searchWord']
        file_path = values['filePath']
        output_filename = values['OutputfilePath']
        words = search_word.split(";")
        words = [word.strip() for word in words]
        sentences = {}

        if file_path.lower().endswith('.pdf'):
            sentences = extract_sentences_with_word_from_pdf(file_path, words)
        elif file_path.lower().endswith('.docx'):
            sentences = extract_sentences_with_word_from_docx(file_path, words)
        else:
            print("Unsupported file format.")

        if sentences:
            save_relevanttext_to_word_file(sentences, output_filename)
            print('Done, Extracted Relevant Text!!!!')
        else:
            print("No relevant text found for the given word(s) or phrase(s). Please try again.")


    if event == 'Extract Equations':

        print('Extracting Equations............')

        file_path = values['filePath']
        output_filename = values['OutputfilePath1']
        equations = []

        if file_path.lower().endswith('.pdf'):
            equations = extract_equations_from_pdf(file_path)
        elif file_path.lower().endswith('.docx'):
            equations = extract_equations_from_docx(file_path)
        else:
            print("Unsupported file format.")
            
        if equations:
            save_equations_to_word_file(equations, output_filename)
            print('Done, Extracted Equations!!!!')
        else:
            print("No relevant equations found in the given file. Please try again.")


    if event == 'Extract Figure Captions':

        print('Extracting Figure Captions............')

        file_path = values['filePath']
        output_filename = values['OutputfilePath2']
        figure_sentences = []

        if file_path.lower().endswith('.pdf'):
            figure_sentences= extract_figure_sentences_from_pdf(file_path)
        elif file_path.lower().endswith('.docx'):
            figure_sentences = extract_figure_sentences_from_word(file_path)
        else:
            print("Unsupported file format.")
            
        if figure_sentences:
            save_figcaptions_to_word(figure_sentences, output_filename)
            print('Done, Extracted Figure Captions!!!!')
        else:
            print("No Figure Captions found in the given file. Please try again.")


    if event == 'Extract Table captions':

        print('Extracting Table Captions............')

        file_path = values['filePath']
        output_filename = values['OutputfilePath3']
        table_sentences = []

        if file_path.lower().endswith('.pdf'):
            table_sentences= extract_table_sentences_from_pdf(file_path)
        elif file_path.lower().endswith('.docx'):
            table_sentences = extract_table_sentences_from_word(file_path)
        else:
            print("Unsupported file format.")
            
        if table_sentences:
            save_tablecaptions_to_word(table_sentences, output_filename)
            print('Done, Extracted Table Captions!!!!')
        else:
            print("No Table Captions found in the given file. Please try again.")


    if event == 'Extract Images':

        print('Extracting Images............')

        file_path = values['filePath']
        output_foldername = values['Outputfolderpath']
        extract_images = []

        if file_path.lower().endswith('.pdf'):
            extract_images= extract_images_from_pdf(file_path)
        elif file_path.lower().endswith('.docx'):
            extract_images = extract_images_from_word(file_path)
        else:
            print("Unsupported file format.")
            
        if extract_images:
            save_images_to_folders(extract_images, output_foldername)
            print('Done, Extracted Images!!!!')
        else:
            print("No Images found in the given file. Please try again.")

    if event == sg.WIN_CLOSED or event == 'Exit':
        break

window.close()