import locale
import PySimpleGUI as sg
import docx 
from openpyxl import Workbook  
from openpyxl.styles import Font  
from openpyxl.styles.borders import Border, Side 
import tabula
import pandas as pd 
import PyPDF2
import re
import docx2txt


##################################################################


def extract_sentences_with_word_from_pdf(filename, word):
    sentences = []
    with open(filename, "rb") as file:
        reader = PyPDF2.PdfReader(file)
        num_pages = len(reader.pages)

        for page_num in range(num_pages):
            page = reader.pages[page_num]
            text = page.extract_text()
            found_sentences = re.findall(r"([^.!?]*{}[^.!?]*[.!?])".format(word), text)
            
            for i, sentence in enumerate(found_sentences):
                if i > 0:
                    sentences.append(found_sentences[i-1]) 
                sentences.append(sentence) 
                if i < len(found_sentences) - 1:
                    sentences.append(found_sentences[i+1]) 

    return sentences

def extract_sentences_with_word_from_docx(filename, word):
    sentences = []
    text = docx2txt.process(filename)
    found_sentences = re.findall(r"([^.!?]*{}[^.!?]*[.!?])".format(word), text)
    
    for i, sentence in enumerate(found_sentences):
        if i > 0:
            sentences.append(found_sentences[i-1]) 
        sentences.append(sentence) 
        if i < len(found_sentences) - 1:
            sentences.append(found_sentences[i+1]) 

    return sentences

search_word = input("Enter the word to search for: ")
file_path = input("Enter the path of the file (PDF or DOCX): ")
sentences = []

if file_path.lower().endswith('.pdf'):
    sentences = extract_sentences_with_word_from_pdf(file_path, search_word)
elif file_path.lower().endswith('.docx'):
    sentences = extract_sentences_with_word_from_docx(file_path, search_word)
else:
    print("Unsupported file format.")

if sentences:
    print("Text relevant to the word '{}':".format(search_word))
    for sentence in sentences:
        print(sentence)
else:
    print("No text found relevant to the word '{}'. Please try another word".format(search_word))


##################################################################


def extract_equations_from_pdf(pdf_path):
    equations = []
    
    with open(pdf_path, 'rb') as file:
        pdf_reader = PyPDF2.PdfReader(file)
        num_pages = len(pdf_reader.pages)
        
        for page_number in range(num_pages):
            page = pdf_reader.pages[page_number]
            text = page.extract_text()
            equation_pattern = r'[^\n=]+=[^\n]+'
            matches = re.findall(equation_pattern, text)
            equations.extend(matches)
    
    return equations

def extract_equations_from_docx(docx_path):
    equations = []

    text = docx2txt.process(docx_path)
    equation_pattern = r'[^\n=]+=[^\n]+'
    matches = re.findall(equation_pattern, text)
    equations.extend(matches)
    
    return equations

file_path = input("Enter the path of the file (PDF or DOCX): ")
equations = []
if file_path.lower().endswith('.pdf'):
    equations = extract_equations_from_pdf(file_path)
    print("The equations in the PDF document are:")
elif file_path.lower().endswith('.docx'):
    equations = extract_equations_from_docx(file_path)
    print("The equations in the Word document are:")
else:
    print("Unsupported file format.")

for equation in equations:
    print(equation)


##################################################################


sg.theme('Purple')
SYMBOL_UP = '▲'
SYMBOL_DOWN = '▼'

def collapse(sectionLayout, visible, key):
    return sg.pin(sg.Column(sectionLayout, visible=visible, key=key))

# LAYOUT
layout = [

          [sg.Text('Thousand & Decimal separator in Word document:', font=(None, 9, 'bold'))],
          [sg.Radio('1.234.567,999', 'RADIO1', default=True, size=(19, 1), key='locale'),
          sg.Radio('1,234,567.999', 'RADIO1', default=False, size=(19, 1))],

          [sg.Text('Source for Word document:', size=(25, 1), font=(None, 9, 'bold')),
           sg.InputText(key='wordDocument'), sg.FileBrowse()],

           [sg.Text('Source for PDF document:', size=(25, 1), font=(None, 9, 'bold')),
            sg.InputText(key='pdfDocument'), sg.FileBrowse()],
           
          [sg.Text('Source of Excel file:', size=(25, 1), font=(None, 9, 'bold'))],
          [sg.Radio('Create a blank Excel file', 'RADIO0', default=False, key='newWb')],

          [sg.Text('Choose any of the following options:', font=(None, 9, 'bold'))],
          [sg.Checkbox('Apply cell border (All Borders)', default=False, key='border')],
          [sg.Checkbox('Copy text format (Bold, Italic & Underline)', default=False, key='formatText')],
          [sg.Checkbox('Convert cell "-" to 0', default=False, key='hyphenToZero')],
          

          [sg.Text('Choose folder to save the word tables:', size=(25, 1), font=(None, 9, 'bold')),
           sg.InputText(key='savedFolder1'), sg.FolderBrowse()],
          [sg.Text('Enter file name (Ex: file.xlsx):', size=(25, 1), font=(None, 9, 'bold')),
           sg.InputText(default_text='file1.xlsx', key='savedName1')],

           [sg.Text('Choose folder to save the pdf tables:', size=(25, 1), font=(None, 9, 'bold')),
            sg.InputText(key='savedFolder2'), sg.FolderBrowse()],
            [sg.Text('Enter file name (Ex: file.xlsx):', size=(25, 1), font=(None, 9, 'bold')),
            sg.InputText(default_text='file2.xlsx', key='savedName2')],

          [sg.Button('Submit'), sg.Button('Exit')],
          [sg.Output(size=(80, 3))]
          ]


##################################################################


def wordToExcel(dict):

    if dict['newWb']:
        wb = Workbook()

    doc = docx.Document(dict['wordDocument'])

    if dict['locale']:
        locale.setlocale(locale.LC_NUMERIC, "en_DK.UTF-8") 
    else:
        locale.setlocale(locale.LC_NUMERIC, "en_US.UTF-8") 

    thin_border = Border(left=Side(style='thin'),
                         right=Side(style='thin'),
                         top=Side(style='thin'),
                         bottom=Side(style='thin'))


    for t in range(len(doc.tables)):
        sheet = wb.create_sheet(title=f'Table {t+1}')
        tableDoc = doc.tables[t]


        rowNumDoc = len(tableDoc.rows)
        colNumDoc = len(tableDoc.columns)

        checkMerge = []

        for i in range(rowNumDoc):
            for j in range(colNumDoc):
 
                excelCell = sheet.cell(row=i + 1, column=j + 1)

                if tableDoc.cell(i, j)._tc not in checkMerge:
                    try:
                        excelCell.value = locale.atoi(
                            tableDoc.cell(i, j).text.replace('(', '-').replace(')', ''))  
                    except ValueError:
                        try:
                            excelCell.value = locale.atof(
                                tableDoc.cell(i, j).text.replace('(', '-').replace(')', ''))  
                            
                        except ValueError:
                            if dict['hyphenToZero'] and tableDoc.cell(i, j).text == "-":
                                excelCell.value = 0
                            else:
                                excelCell.value = tableDoc.cell(i, j).text  
                            

                checkMerge.append(tableDoc.cell(i, j)._tc)


                if dict['border']:
                    excelCell.border = thin_border  
                if dict['formatText']:
                    if tableDoc.cell(i, j).text:  
                        try:
                            checkBold = tableDoc.cell(i, j).paragraphs[0].runs[
                                0].bold  
                            checkItalic = tableDoc.cell(i, j).paragraphs[0].runs[
                                0].italic  
                            checkUnderline = tableDoc.cell(i, j).paragraphs[0].runs[
                                0].underline  
                        except:
                            continue


                        excelCell.font = Font(bold=checkBold, italic=checkItalic,
                                              underline=('single' if checkUnderline else None))
                    excelCell.border = thin_border


    if not dict['savedName1'].endswith('.xlsx'):
        dict['savedName1'] += '.xlsx'
    wb.save((dict['savedFolder1'] + '/' + dict['savedName1']))


##################################################################


def pdfToExcel(dict):
    if dict['newWb']:
        wb = Workbook()
    tables = tabula.read_pdf(dict['pdfDocument'], pages='all', multiple_tables=True)

    if dict['locale']:
        locale.setlocale(locale.LC_NUMERIC, "en_DK.UTF-8")  
    else:
        locale.setlocale(locale.LC_NUMERIC, "en_US.UTF-8")  

    
    thin_border = Border(left=Side(style='thin'),
                         right=Side(style='thin'),
                         top=Side(style='thin'),
                         bottom=Side(style='thin'))

    
    for table_num, table in enumerate(tables):
        
        sheet = wb.create_sheet(title=f'Table {table_num + 1}')
        excel_writer = pd.ExcelWriter((dict['savedFolder2'] + '/' + dict['savedName2']), engine='xlsxwriter')

    for i, table in enumerate(tables):
    
        sheet_name = f"Table {i+1}"
        table.to_excel(excel_writer, sheet_name=sheet_name, index=False)

    excel_writer.save()

        
    for i, row in enumerate(tables):
            
            for j, cell in enumerate(row):
                
                excelCell = sheet.cell(row=i + 1, column=j + 1)

                
                try:
                    excelCell.value = locale.atof(cell.replace('(', '-').replace(')', ''))
                except ValueError:
                    if dict['hyphenToZero'] and cell == "-":
                        excelCell.value = 0
                    else:
                        excelCell.value = cell

                
                if dict['border']:
                    excelCell.border = thin_border

                
                if dict['formatText']:
                    excelCell.font = Font(bold=False,
                                          italic=False,
                                          underline=False)

    
    if not dict['savedName2'].endswith('.xlsx'):
        dict['savedName2'] += '.xlsx'
    wb.save((dict['savedFolder2'] + '/' + dict['savedName2']))

window = sg.Window('Copy Word Tables or PDF Tables to Excel', layout)

while True:  
    event, values = window.read()

    if event == sg.WIN_CLOSED or event == 'Exit':
        break
    if event == 'Submit':

        print('Running............')

        keys = ['newWb', 'wordDocument','pdfDocument',
                'locale',
                'border','hyphenToZero', 'formatText',
                'savedFolder1', 'savedName1',
                'savedFolder2', 'savedName2'
                
                ]
        convertDict = {x: window[x].get() for x in keys}

        if convertDict['wordDocument'] and convertDict['savedFolder1']:
            wordToExcel(convertDict)

        if convertDict['pdfDocument'] and convertDict['savedFolder2']:
            pdfToExcel(convertDict)
        print('Done!!!!')

    if event.startswith('-OPEN SEC1-'):
        opened1 = not opened1
        window['-OPEN SEC1-'].update(SYMBOL_UP if opened1 else SYMBOL_DOWN)
        window['-SEC1-'].update(visible=opened1)

window.close()


##################################################################