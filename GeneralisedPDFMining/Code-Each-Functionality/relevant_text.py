import PyPDF2
import docx
import docx2txt
import re

def save_sentences_to_word_file(sentences, output_filename):
    doc = docx.Document()
    for word, sentences in sentences.items():
        doc.add_paragraph(f"The relevant text for the word '{word}' is:")
        for sentence in sentences:
            doc.add_paragraph(sentence)
        doc.add_paragraph("\n")
    doc.save(output_filename)

def extract_sentences_with_word_from_pdf(filename, words):
    sentences = {word: [] for word in words}
    previous_sentence = None
    sentence_set = set()
    stop_extraction = False

    with open(filename, "rb") as file:
        reader = PyPDF2.PdfReader(file)
        num_pages = len(reader.pages)

        for page_num in range(num_pages):
            page = reader.pages[page_num]
            text = page.extract_text()

            if stop_extraction:
                break

            for word in words:
                if " " in word:
                    found_sentences = re.findall(r"([^.!?]*{}[^.!?]*[.!?])".format(re.escape(word)), text)
                else:
                    found_sentences = re.findall(r"([^.!?]*\b{}\b[^.!?]*[.!?])".format(re.escape(word)), text)

                for sentence in found_sentences:
                    if sentence in sentence_set:
                        continue  # Skip repeated sentence

                    if previous_sentence is not None and previous_sentence == sentence:
                        continue  # Skip repeated sentence

                    if sentence.startswith("Figure") or sentence.startswith("Page"):
                        continue  # Skip sentences starting with "Figure" or "Page"

                    if "references" in sentence.lower():
                        stop_extraction = True
                        break

                    sentences[word].append(sentence)
                    sentence_set.add(sentence)
                    previous_sentence = sentence

    return sentences

def extract_sentences_with_word_from_docx(filename, words):
    sentences = {word: [] for word in words}
    previous_sentence = None
    sentence_set = set()
    stop_extraction = False

    text = docx2txt.process(filename)

    for word in words:
        if " " in word:
            found_sentences = re.findall(r"([^.!?]*{}[^.!?]*[.!?])".format(re.escape(word)), text)
        else:
            found_sentences = re.findall(r"([^.!?]*\b{}\b[^.!?]*[.!?])".format(re.escape(word)), text)

        for sentence in found_sentences:
            if sentence in sentence_set:
                continue  # Skip repeated sentence

            if previous_sentence is not None and previous_sentence == sentence:
                continue  # Skip repeated sentence

            if sentence.startswith("Figure") or sentence.startswith("Page"):
                continue  # Skip sentences starting with "Figure" or "Page"

            if "references" in sentence.lower():
                stop_extraction = True
                break

            sentences[word].append(sentence)
            sentence_set.add(sentence)
            previous_sentence = sentence

    return sentences

search_input = input("Enter the word(s) or phrase(s) to search for (separated by a semicolon): ")
file_path = input("Enter the path of the file (PDF or DOCX): ")
output_filename = "output.docx"

words = search_input.split(";")
words = [word.strip() for word in words]
sentences = {}

if file_path.lower().endswith('.pdf'):
    sentences = extract_sentences_with_word_from_pdf(file_path, words)
elif file_path.lower().endswith('.docx'):
    sentences = extract_sentences_with_word_from_docx(file_path, words)
else:
    print("Unsupported file format.")

if sentences:
    save_sentences_to_word_file(sentences, output_filename)
else:
    print("No relevant text found for the given word(s) or phrase(s). Please try again.")


